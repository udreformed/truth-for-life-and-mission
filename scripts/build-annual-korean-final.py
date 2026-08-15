#!/usr/bin/env python3
"""Build the detailed Korean 365-day Bible-reading review DOCX."""

import calendar
import glob
import json
import re
import sys
from collections import Counter
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


NT_BOOKS = {
    "마", "막", "눅", "요", "행", "롬", "고전", "고후", "갈", "엡", "빌", "골",
    "살전", "살후", "딤전", "딤후", "딛", "몬", "히", "약", "벧전", "벧후",
    "요일", "요이", "요삼", "유", "계",
}

BOOK_NAMES = {
    "창":"창세기", "출":"출애굽기", "레":"레위기", "민":"민수기", "신":"신명기",
    "수":"여호수아", "삿":"사사기", "룻":"룻기", "삼상":"사무엘상", "삼하":"사무엘하",
    "왕상":"열왕기상", "왕하":"열왕기하", "대상":"역대상", "대하":"역대하", "스":"에스라",
    "느":"느헤미야", "에":"에스더", "욥":"욥기", "시":"시편", "잠":"잠언", "전":"전도서",
    "아":"아가", "사":"이사야", "렘":"예레미야", "애":"예레미야애가", "겔":"에스겔",
    "단":"다니엘", "호":"호세아", "욜":"요엘", "암":"아모스", "옵":"오바댜", "오":"오바댜",
    "욘":"요나", "미":"미가", "나":"나훔", "합":"하박국", "습":"스바냐", "학":"학개",
    "슥":"스가랴", "말":"말라기", "마":"마태복음", "막":"마가복음", "눅":"누가복음",
    "요":"요한복음", "행":"사도행전", "롬":"로마서", "고전":"고린도전서", "고후":"고린도후서",
    "갈":"갈라디아서", "엡":"에베소서", "빌":"빌립보서", "골":"골로새서",
    "살전":"데살로니가전서", "살후":"데살로니가후서", "딤전":"디모데전서", "딤후":"디모데후서",
    "딛":"디도서", "몬":"빌레몬서", "히":"히브리서", "약":"야고보서", "벧전":"베드로전서",
    "벧후":"베드로후서", "요일":"요한일서", "요이":"요한이서", "요삼":"요한삼서",
    "유":"유다서", "계":"요한계시록",
}


def set_font(run, size=9.3, bold=False, color=None):
    # Use the TrueType Apple Gothic face rather than the TTC-only system face.
    # This keeps Korean glyphs available in both Word and LibreOffice PDF previews.
    name = "AppleGothic"
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend((fld_char1, instr, fld_char2))
    set_font(run, size=8, color="69756F")


def expand_passage(text):
    # Supports the one multi-book entry: 암 9,오 1.
    parts = [p.strip() for p in text.split(",")]
    expanded = []
    current = None
    for part in parts:
        m = re.match(r"^([가-힣]+)?\s*(.+)$", part)
        if not m:
            expanded.append(part)
            continue
        book, ref = m.groups()
        if book:
            current = book
        if current in BOOK_NAMES:
            suffix = "" if re.search(r"[:;]", ref) else "장"
            normalized_ref = ref.replace(";", ":").replace("~", "–")
            normalized_ref = re.sub(r"(?<=\d)-(?=\d)", "–", normalized_ref)
            expanded.append(f"{BOOK_NAMES[current]} {normalized_ref}{suffix}")
        else:
            expanded.append(part)
    return " / ".join(expanded)


def normalize_plan(plan):
    for item in plan:
        if item["month"] == 2 and item["day"] == 28:
            item["oldTestament"] = "레 9~10"
            item["newTestament"] = "눅 14~15"
        if item["month"] == 9 and item["day"] <= 16:
            item["oldTestament"] = f"욥 {item['oldTestament']}"
    return plan


def validate(plan, drafts):
    expected_keys = [
        (month, day)
        for month in range(1, 13)
        for day in range(1, calendar.monthrange(2026, month)[1] + 1)
    ]
    draft_by_key = {(x["month"], x["day"]): x for x in drafts}
    problems = []
    if len(drafts) != 365 or len(draft_by_key) != 365:
        problems.append(f"원고 개수/고유 날짜 오류: {len(drafts)}/{len(draft_by_key)}")
    if set(expected_keys) != set(draft_by_key):
        problems.append("365일 날짜 집합이 달력과 일치하지 않음")
    if len(plan) != 365:
        problems.append(f"진도표 개수 오류: {len(plan)}")
    all_texts = []
    sentence_counts = Counter()
    lengths = []
    for key in expected_keys:
        draft = draft_by_key.get(key)
        if not draft:
            continue
        for field in ("first", "second"):
            text = draft.get(field, "").strip()
            count = len(re.findall(r"[.!?](?:\s|$)", text))
            sentence_counts[count] += 1
            lengths.append(len(text))
            all_texts.append(text)
            if count not in (2, 3):
                problems.append(f"{key} {field} 문장 수 {count}")
    duplicates = [t for t, count in Counter(all_texts).items() if count > 1]
    if duplicates:
        problems.append(f"완전 중복 원고 {len(duplicates)}개")
    if problems:
        raise ValueError("; ".join(problems))
    return {
        "days": len(draft_by_key),
        "sections": len(all_texts),
        "sentenceCounts": dict(sentence_counts),
        "exactDuplicates": len(duplicates),
        "minChars": min(lengths),
        "maxChars": max(lengths),
        "avgChars": round(sum(lengths) / len(lengths), 1),
    }


def add_heading(doc, text, size=15):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(6)
    set_font(p.add_run(text), size=size, bold=True, color="164C3C")
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.12
    if bold_prefix and text.startswith(bold_prefix):
        set_font(p.add_run(bold_prefix), size=9.5, bold=True)
        set_font(p.add_run(text[len(bold_prefix):]), size=9.5)
    else:
        set_font(p.add_run(text), size=9.5)
    return p


def build(plan_path, draft_glob, output_path):
    plan = normalize_plan(json.loads(Path(plan_path).read_text(encoding="utf-8")))
    drafts = []
    for path in sorted(glob.glob(draft_glob)):
        drafts.extend(json.loads(Path(path).read_text(encoding="utf-8")))
    metrics = validate(plan, drafts)
    draft_by_key = {(x["month"], x["day"]): x for x in drafts}

    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(1.45)
    sec.bottom_margin = Cm(1.45)
    sec.left_margin = Cm(1.55)
    sec.right_margin = Cm(1.55)
    add_page_number(sec.footer.paragraphs[0])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(42)
    set_font(p.add_run("365일 성경읽기"), size=24, bold=True, color="164C3C")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("읽기 전 핵심 내용 안내"), size=19, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    set_font(p.add_run("출판 보완 검토본 · 평년용 365일 진도"), size=11, color="7A5A20")

    doc.add_paragraph()
    add_body(doc, "365일 동안 매일 약 15분 분량의 구약과 신약 본문을 읽으며, 1년에 성경 전체를 통독하도록 계획된 성경읽기표입니다.")
    add_body(doc, "신약 통독을 마친 뒤에는 구약 역사·예언서와 시편·잠언·전도서·아가를 함께 읽습니다. 각 설명은 성경을 펼치기 전에 그날 만나게 될 인물, 사건, 명령, 논증과 장면의 전환을 미리 파악하도록 작성했습니다.")
    add_body(doc, "평년 2월 28일에는 레위기 9–10장과 누가복음 14–15장을 합쳐 읽습니다. 윤년에는 같은 분량을 2월 28일과 29일로 나누어 적용합니다.")

    add_heading(doc, "작성 원칙", size=14)
    for text in [
        "1. 두 문장 모두 ‘무슨 내용이 나오는가’를 먼저 알립니다. 사건 본문은 인물과 행동과 결과를, 율법은 규례의 대상과 절차를, 시와 지혜문학은 화자와 중심 대조를, 서신은 논증과 명령의 순서를 적었습니다.",
        "2. 본문에 직접 나오지 않는 평가나 오늘의 적용을 앞세우지 않았습니다. 신학적 의미는 본문 자체가 분명히 선언하거나, 다음 사건을 이해하는 데 꼭 필요한 범위에서만 짧게 남겼습니다.",
        "3. 구약 본문을 곧바로 그리스도 예표로 결론짓던 문장은 해당 장의 실제 제도·행동·대화·결말로 바꾸었습니다. 본문 자체에 메시아 약속이 명시된 경우에는 그 약속을 내용으로 소개했습니다.",
        "4. 욥기의 친구들과 엘리후의 말은 성경의 최종 판정처럼 단정하지 않고 ‘누가 무엇을 주장하는가’를 구별했습니다. 예언서와 요한계시록도 상징을 특정 도식으로 확정하기보다 실제 환상과 선언의 순서를 안내했습니다.",
        "5. 기본은 두 문장이지만 세 장 이상을 읽어 중요한 장면 전환이 빠질 수 있는 날은 세 문장까지 허용했습니다. 분량을 기계적으로 통일하기보다 실제 읽기 범위를 충실히 안내하는 것을 우선했습니다.",
        "6. 성경과 참고 자료의 문장을 옮기지 않고, 읽기 전 안내에 맞는 한국어 문장으로 새롭게 요약했습니다.",
    ]:
        add_body(doc, text)

    add_heading(doc, "장르별 편집자 주", size=14)
    add_body(doc, "아가: 화자 표시는 성경 원문에 직접 붙어 있는 표제가 아닙니다. 이 원고는 사용자의 『성경 66권 해설서』가 채택한 트렘퍼 롱맨 3세의 23편 구분을 주된 기준으로 삼아 남자·여인·예루살렘의 딸들·형제들의 목소리를 구별하며, 화자가 확실하지 않은 곳은 단정하지 않습니다.")
    add_body(doc, "요한계시록: 이 환상들은 당시 일곱 교회에 주어진 말씀이면서 교회 시대와 최종 완성을 함께 바라봅니다. 이 원고는 상징을 특정 현대 사건과 일대일로 연결하지 않고, 하늘 성전과 인·나팔·대접, 용과 두 짐승, 바벨론, 최후 심판과 새 창조의 문학적 흐름 및 구약적 배경을 중심으로 안내합니다.")

    doc.add_section(WD_SECTION.NEW_PAGE)
    current_month = None
    for item in plan:
        key = (item["month"], item["day"])
        draft = draft_by_key[key]
        if item["month"] != current_month:
            if current_month is not None:
                doc.add_section(WD_SECTION.NEW_PAGE)
            current_month = item["month"]
            add_heading(doc, f"{current_month}월", size=17)

        table = doc.add_table(rows=4, cols=2)
        table.style = "Table Grid"
        table.autofit = False
        table.columns[0].width = Cm(2.25)
        table.columns[1].width = Cm(15.4)
        for row in table.rows:
            for cell in row.cells:
                set_cell_margins(cell)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        date_cell = table.cell(0, 0)
        date_cell.merge(table.cell(0, 1))
        set_shading(date_cell, "DDEBE5")
        p = date_cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(f"{item['month']}월 {item['day']}일"), size=10.7, bold=True, color="164C3C")

        second_abbr = item["newTestament"].split()[0]
        second_label = "신약 미리보기" if second_abbr in NT_BOOKS else "성문서 미리보기"
        rows = [
            ("읽기 본문", f"{expand_passage(item['oldTestament'])} / {expand_passage(item['newTestament'])}"),
            ("구약 미리보기", draft["first"]),
            (second_label, draft["second"]),
        ]
        for idx, (label, value) in enumerate(rows, 1):
            label_cell = table.cell(idx, 0)
            value_cell = table.cell(idx, 1)
            set_shading(label_cell, "F3F5F2")
            label_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            value_cell.paragraphs[0].paragraph_format.space_after = Pt(0)
            value_cell.paragraphs[0].paragraph_format.line_spacing = 1.04
            set_font(label_cell.paragraphs[0].add_run(label), size=8.7, bold=True, color="3E514A")
            set_font(value_cell.paragraphs[0].add_run(value), size=9.0)
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_after = Pt(1)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "편집 검증 기록", size=17)
    evaluations = [
        ("범위 완성도", f"365일 전체, 본문 설명 {metrics['sections']}개를 수록했습니다. 모든 날짜가 고유하며 빠진 날짜가 없습니다."),
        ("본문 충실도", "730개 설명을 단계별로 다시 읽고 사건·인물·명령·논증과 장면 전환이 먼저 제시되는지 검토했습니다. 여러 검토 단계에서 겹쳐 수정한 항목을 하나로 계산하면 최초 상세 원고의 설명 106개를 실질적으로 다시 썼습니다."),
        ("서술 방향", "개혁주의적 신학 방향은 해설의 배경으로 유지하되, 이번 원고의 전면에는 해석 체계보다 성경 본문 자체의 내용을 놓았습니다. 특히 구약 제도와 역사, 욥의 대화에서 후대의 교리적 결론을 먼저 말하지 않도록 조정했습니다."),
        ("문체", f"각 설명은 2–3문장이고 길이는 {metrics['minChars']}–{metrics['maxChars']}자, 평균 {metrics['avgChars']}자입니다. 완전히 동일한 원고는 {metrics['exactDuplicates']}개입니다. 세 장 이상 범위 가운데 장면 보완이 꼭 필요한 항목만 세 문장으로 늘렸습니다."),
        ("남은 한계", "두 문장이라는 제한 때문에 하루에 여러 장을 읽는 날은 모든 장면을 열거하지 못하고 중심 전환을 선택했습니다. ‘보여 준다’나 ‘가르친다’는 표현도 일부 남아 있지만, 그 문장이 본문에 실제로 제시된 비유·명령·논증을 요약하는 경우에는 내용 안내로 판단해 유지했습니다."),
        ("추가 관찰이 필요한 부분", "왕조의 병행 연대, 선지서의 역사 배경, 히브리어·그리스어의 세부 의미, 시편의 메시아적 적용은 실제 사이트에서 몇 달간 사용하면서 발견되는 어색함이나 과잉 적용을 계속 교정하는 것이 좋습니다."),
        ("종합 판단", "이전 초안보다 읽기 전 안내의 목적이 분명해졌고, 구체적 사건과 논증의 비중이 높아졌습니다. 웹사이트 시험 운영에 사용할 수 있는 재검토본이지만, 실제로 매일 읽으면서 빠진 핵심 장면과 지나치게 압축된 날짜를 표시한 뒤 확정본으로 다듬는 것이 적절합니다."),
    ]
    for label, value in evaluations:
        add_body(doc, f"{label}: {value}", bold_prefix=f"{label}: ")

    add_heading(doc, "참고 방향과 자료", size=15)
    for text in [
        "성경 본문과 사용자가 제공한 365일 진도표를 최우선 기준으로 삼았습니다.",
        "기존 리딩지저스 요약 자료는 성경 전체의 구속사적 구조와 단락 확인을 위한 보조 자료로 참고했습니다.",
        "최용진, 『성경 66권 해설서』 ver. 3.9를 아가의 화자 구분과 요한계시록의 구조·상징 해석 방향을 확인하는 직접 기준으로 사용했습니다.",
        "아가의 화자와 장면 전환은 위 해설서가 채택한 트렘퍼 롱맨 3세의 23편 구분을 따랐으며, 1장 안내에서 이 기준을 명시했습니다.",
        "요한계시록은 위 해설서의 큰 구조와 절충적·상징적 읽기를 따라 하늘 성전, 인·나팔·대접, 용과 두 짐승, 바벨론, 최후 심판, 새 창조의 흐름을 구분했습니다.",
        "R.C. 스프룰이 편집한 『Reformation Study Bible』과 개혁주의의 언약·은혜·그리스도 중심적 성경 이해를 신학적 방향으로 삼았습니다.",
        "NET Bible/Bible.org의 번역·문맥 자료와 BibleRef의 장별 개요는 보조적인 사실 확인 자료로 두며, 이들의 모든 세부 신학 입장을 최종 기준으로 삼지는 않습니다.",
        "참고 자료의 문장을 복사하지 않았으며, 모든 설명은 이 문서를 위해 새로 작성한 요약입니다.",
    ]:
        add_body(doc, text)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(json.dumps({"output": str(output_path), "metrics": metrics}, ensure_ascii=False))


if __name__ == "__main__":
    build(sys.argv[1], sys.argv[2], sys.argv[3])
