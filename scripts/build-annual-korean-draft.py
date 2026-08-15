#!/usr/bin/env python3
"""Build a review-only Korean draft for the 365-day Bible reading plan."""

import json
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BOOKS = {
    "창": ("창세기", "창조와 타락, 언약의 시작과 인간의 실패 속에서도 약속을 이루시는 하나님의 신실하심"),
    "출": ("출애굽기", "구원하신 백성을 자기 백성으로 삼으시고 말씀과 임재로 인도하시는 언약의 하나님"),
    "레": ("레위기", "거룩하신 하나님께 나아가는 속죄의 길과 구원받은 백성에게 요구되는 거룩한 삶"),
    "민": ("민수기", "광야에서 드러난 불신앙에도 불구하고 약속의 땅으로 백성을 이끄시는 하나님의 인내와 신실하심"),
    "신": ("신명기", "언약의 말씀을 기억하고 마음과 삶으로 하나님을 사랑하며 순종하라는 부르심"),
    "수": ("여호수아", "약속의 땅을 주시는 하나님의 신실하심과 언약 백성에게 요구되는 믿음의 순종"),
    "삿": ("사사기", "죄와 압제와 구원이 반복되는 시대에 참된 왕과 온전한 구원자가 필요함"),
    "룻": ("룻기", "평범한 충성과 섭리를 통해 다윗과 그리스도의 계보를 이어 가시는 하나님의 은혜"),
    "삼상": ("사무엘상", "사람의 외모가 아니라 중심을 보시며 자기 뜻에 합한 왕을 세우시는 하나님의 주권"),
    "삼하": ("사무엘하", "다윗 언약과 왕의 실패를 함께 보여 주며 완전한 메시아 왕을 기다리게 하는 역사"),
    "왕상": ("열왕기상", "왕과 백성의 순종 여부에 따라 드러나는 언약의 복과 심판, 그리고 선지자를 통한 하나님의 말씀"),
    "왕하": ("열왕기하", "지속되는 우상숭배와 불순종을 심판하시면서도 언약의 약속을 보존하시는 하나님"),
    "대상": ("역대상", "다윗 왕조와 예배 공동체의 뿌리를 돌아보며 언약과 성전 예배의 중심을 회복하도록 이끄는 말씀"),
    "대하": ("역대하", "성전과 다윗 왕조를 중심으로 회개와 개혁을 촉구하고 포로 이후의 소망을 열어 주는 역사"),
    "스": ("에스라", "포로에서 돌아온 백성이 말씀과 예배를 회복하도록 섭리하시는 하나님"),
    "느": ("느헤미야", "무너진 성벽과 공동체를 재건하고 말씀에 따라 언약을 새롭게 하는 회복"),
    "에": ("에스더", "하나님의 이름이 보이지 않는 순간에도 자기 백성을 보존하시는 숨은 섭리"),
    "욥": ("욥기", "고난을 단순한 인과응보로 설명할 수 없으며 피조물의 한계를 넘어 주권적으로 일하시는 하나님"),
    "시": ("시편", "기쁨과 탄식과 회개와 소망을 하나님께 아뢰며 언약의 왕과 하나님의 통치를 바라보는 믿음"),
    "잠": ("잠언", "여호와를 경외하는 데서 시작하여 말과 관계와 일상의 선택으로 나타나는 지혜"),
    "전": ("전도서", "해 아래 삶의 헛됨을 직시하면서 창조주를 경외하고 그분이 주시는 선물을 감사히 받는 지혜"),
    "아": ("아가", "언약 안에서 주어진 사랑의 아름다움과 배타적 충실함을 노래하는 말씀"),
    "사": ("이사야", "거룩하신 하나님의 심판과 남은 자의 회복, 고난받는 종을 통해 임할 구원"),
    "렘": ("예레미야", "완고한 죄를 심판하시면서도 마음에 율법을 새기는 새 언약을 약속하시는 하나님"),
    "애": ("예레미야애가", "예루살렘의 멸망을 슬퍼하면서도 하나님의 긍휼과 신실하심을 붙드는 탄식"),
    "겔": ("에스겔", "떠난 영광과 공의로운 심판, 새 마음과 새 영과 회복된 성전을 약속하시는 하나님"),
    "단": ("다니엘", "세상 제국 위에서 통치하시며 환난 속에서도 자기 백성과 영원한 나라를 보존하시는 하나님"),
    "호": ("호세아", "배신한 언약 백성을 심판하면서도 신실한 사랑으로 다시 부르시는 하나님"),
    "욜": ("요엘", "여호와의 날을 앞두고 회개를 촉구하며 모든 백성에게 부어질 성령을 약속하는 말씀"),
    "암": ("아모스", "형식적인 예배와 사회적 불의를 심판하고 정의와 공의를 요구하시는 하나님"),
    "옵": ("오바댜", "교만과 형제에 대한 폭력을 심판하시고 자기 나라를 세우시는 하나님"),
    "욘": ("요나", "불순종하는 선지자와 원수 된 민족에게까지 긍휼을 베푸시는 선교의 하나님"),
    "미": ("미가", "죄를 심판하시되 베들레헴의 통치자와 남은 자의 회복을 약속하시는 하나님"),
    "나": ("나훔", "폭력적인 제국을 심판하고 압제받는 백성의 피난처가 되시는 하나님"),
    "합": ("하박국", "악이 성행하는 현실 속에서도 하나님의 주권을 신뢰하며 믿음으로 사는 길"),
    "습": ("스바냐", "여호와의 날의 심판과 정결하게 된 남은 백성을 향한 회복의 기쁨"),
    "학": ("학개", "무너진 예배를 회복하고 하나님의 임재와 장차 올 영광을 우선하도록 부르시는 말씀"),
    "슥": ("스가랴", "성전 재건을 넘어 오실 왕과 목자와 정결의 샘을 바라보게 하는 회복의 환상"),
    "말": ("말라기", "무너진 예배와 언약적 신실함을 책망하고 주의 길을 준비할 사자를 약속하시는 말씀"),
    "마": ("마태복음", "구약의 약속을 성취하고 하나님 나라를 선포하시는 왕이신 예수 그리스도"),
    "막": ("마가복음", "섬기며 자기 목숨을 대속물로 내어 주시는 하나님의 아들 예수 그리스도"),
    "눅": ("누가복음", "성령으로 사역하시며 잃어버린 자와 주변부의 사람들에게 구원을 베푸시는 인자 예수"),
    "요": ("요한복음", "말씀이 육신이 되어 오신 하나님의 아들을 믿고 생명을 얻도록 증언하는 복음"),
    "행": ("사도행전", "부활하신 그리스도께서 성령을 통해 교회를 세우고 복음을 땅끝까지 전하게 하시는 역사"),
    "롬": ("로마서", "모든 사람의 죄와 하나님의 의, 믿음으로 얻는 칭의와 복음에 합당한 새 삶"),
    "고전": ("고린도전서", "십자가의 지혜로 분열과 세속성을 바로잡고 거룩한 교회 공동체를 세우는 말씀"),
    "고후": ("고린도후서", "약함 속에서 나타나는 그리스도의 능력과 새 언약의 일꾼이 감당하는 화해의 직분"),
    "갈": ("갈라디아서", "율법의 행위가 아니라 그리스도를 믿음으로 의롭다 함을 받고 성령으로 사는 자유"),
    "엡": ("에베소서", "그리스도 안에서 베푸신 구원의 은혜와 하나 된 교회가 살아갈 새 삶"),
    "빌": ("빌립보서", "복음에 합당하게 살며 낮아지신 그리스도를 본받아 기쁨으로 함께 전진하는 삶"),
    "골": ("골로새서", "만물과 교회의 머리이신 그리스도의 충만함과 그분 안에서 살아가는 새 사람"),
    "살전": ("데살로니가전서", "재림의 소망 가운데 거룩과 사랑과 인내로 살아가도록 격려하는 말씀"),
    "살후": ("데살로니가후서", "주의 날에 관한 혼란을 바로잡고 박해 속에서 흔들리지 않도록 권면하는 말씀"),
    "딤전": ("디모데전서", "복음의 진리를 지키며 바른 예배와 지도력으로 하나님의 집을 세우는 원리"),
    "딤후": ("디모데후서", "고난 속에서도 복음을 부끄러워하지 않고 말씀을 끝까지 맡아 전하라는 유언적 권면"),
    "딛": ("디도서", "바른 교훈이 선한 행실과 건강한 교회 질서로 열매 맺어야 함을 가르치는 말씀"),
    "몬": ("빌레몬서", "복음 안에서 종과 주인의 관계를 형제의 관계로 새롭게 하는 용서와 화해"),
    "히": ("히브리서", "옛 언약의 모든 모형을 성취하신 완전한 대제사장 예수와 새 언약의 우월함"),
    "약": ("야고보서", "행함으로 진실성이 드러나는 믿음과 시험 속에서 필요한 하늘의 지혜"),
    "벧전": ("베드로전서", "나그네 된 성도가 고난 속에서도 산 소망과 거룩함으로 살아가는 길"),
    "벧후": ("베드로후서", "거짓 교훈을 분별하고 그리스도의 재림을 기다리며 은혜와 지식에서 자라는 삶"),
    "요일": ("요한일서", "성육신하신 그리스도를 믿고 빛과 사랑 가운데 행하는 자에게 있는 구원의 확신"),
    "요이": ("요한이서", "진리와 사랑 안에 거하면서 그리스도에 관한 거짓 가르침을 경계하는 말씀"),
    "요삼": ("요한삼서", "복음을 위해 수고하는 일꾼을 환대하고 교만한 권력욕을 경계하는 말씀"),
    "유": ("유다서", "교회에 침투한 거짓 교사에 맞서 단번에 주신 믿음을 위해 힘써 싸우라는 권면"),
    "계": ("요한계시록", "죽임당하셨으나 승리하신 어린양의 통치와 교회의 인내, 최후 심판과 새 창조의 소망"),
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_font(run, name="Apple SD Gothic Neo", size=9, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def parse_passage(text):
    m = re.match(r"([가-힣]+)\s*(.+)", text.strip())
    if not m or m.group(1) not in BOOKS:
        return text.strip(), text.strip(), "본문의 문맥과 구속사적 위치를 중심으로 읽어야 하는 말씀", []
    abbr, chapters = m.groups()
    name, theme = BOOKS[abbr]
    chapter_nums = [int(n) for n in re.findall(r"\d+", chapters)]
    if len(chapter_nums) >= 2 and ("~" in chapters or "-" in chapters):
        chapter_nums = list(range(chapter_nums[0], chapter_nums[1] + 1))
    elif chapter_nums:
        chapter_nums = [chapter_nums[0]]
    return f"{name} {chapters}장", name, theme, chapter_nums


def load_topics(source_path):
    text = Path(source_path).read_text(encoding="utf-8")
    raw = text.split("export const readingSummaryData =", 1)[1].split(" as const satisfies", 1)[0].strip()
    data = json.loads(raw)
    topics = {}
    full_names = {name for name, _ in BOOKS.values()}
    for week in data.values():
        for day in week.values():
            title = day.get("ko", {}).get("title", "")
            match = re.match(r"([가-힣]+)\s+(\d+)(?:\s*[~\-–]\s*(\d+))?장?\s*(.*)", title)
            if not match or match.group(1) not in full_names:
                continue
            book, start, end, topic = match.groups()
            start, end = int(start), int(end or start)
            topic = topic.strip(" :,-–")
            if not topic:
                continue
            for chapter in range(start, end + 1):
                topics[(book, chapter)] = topic
    return topics


def make_summary(passage, topics):
    display, book, theme, chapters = parse_passage(passage)
    found = []
    for chapter in chapters:
        topic = topics.get((book, chapter))
        if topic and topic not in found:
            found.append(topic)
    if found:
        focus = " 그리고 ".join(found[:2])
        sentence1 = f"{display}의 핵심 주제는 ‘{focus}’입니다."
    else:
        sentence1 = f"{display}은 다음 흐름 안에서 읽습니다: {theme}."
    sentence2 = (
        "본문의 사건과 가르침을 인간의 공로나 도덕적 모범에만 가두지 말고, "
        "언약에 신실하신 하나님의 주권과 죄인을 구원하시는 은혜의 관점에서 살펴봅니다."
    )
    sentence3 = (
        "또한 이 말씀이 그리스도 안에서 어떻게 성취되거나 복음에 합당한 믿음과 순종으로 이어지는지 묵상합니다."
    )
    return " ".join((sentence1, sentence2, sentence3))


def classify_second(passage):
    abbr = passage.split()[0]
    return "신약" if abbr in {"마","막","눅","요","행","롬","고전","고후","갈","엡","빌","골","살전","살후","딤전","딤후","딛","몬","히","약","벧전","벧후","요일","요이","요삼","유","계"} else "성문서"


def build(plan_path, source_path, output_path):
    plan = json.loads(Path(plan_path).read_text(encoding="utf-8"))
    # The supplied plan is a 365-day plan. In a common year, Feb 28 combines
    # the two readings that are separated into Feb 28 and 29 in a leap year.
    for item in plan:
        if item["month"] == 2 and item["day"] == 28:
            item["oldTestament"] = "레 9~10"
            item["newTestament"] = "눅 14~15"
    topics = load_topics(source_path)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Cm(1.5)
    sec.bottom_margin = Cm(1.5)
    sec.left_margin = Cm(1.7)
    sec.right_margin = Cm(1.7)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(title.add_run("365일 성경읽기 한국어 본문 설명 — 1차 초안"), size=18, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("검토용 · 사이트 게시 전 수정 필요"), size=10, bold=True)

    for text in [
        "365일 동안 매일 약 15분 분량의 구약과 신약 본문을 읽으며, 1년에 성경 전체를 통독하도록 계획된 성경읽기표입니다.",
        "이 문서는 전체 365일 구조와 개혁주의적 문체 방향을 확인하기 위한 1차 초안입니다. 날짜와 본문 범위는 제공된 수정본 진도표에서 추출했습니다.",
    ]:
        p = doc.add_paragraph()
        set_font(p.add_run(text), size=10)

    current_month = None
    for item in plan:
        if item["month"] != current_month:
            if current_month is not None:
                doc.add_section(WD_SECTION.NEW_PAGE)
            current_month = item["month"]
            h = doc.add_paragraph()
            set_font(h.add_run(f"{current_month}월"), size=15, bold=True)

        table = doc.add_table(rows=4, cols=2)
        table.autofit = False
        table.columns[0].width = Cm(2.4)
        table.columns[1].width = Cm(15.2)
        date_cell = table.cell(0, 0)
        date_cell.merge(table.cell(0, 1))
        set_cell_shading(date_cell, "E8F1ED")
        set_font(date_cell.paragraphs[0].add_run(f'{item["month"]}월 {item["day"]}일'), size=11, bold=True)

        second_kind = classify_second(item["newTestament"])
        rows = [
            ("읽기 본문", f'{item["oldTestament"]} / {item["newTestament"]}'),
            ("구약 설명", make_summary(item["oldTestament"], topics)),
            (f"{second_kind} 설명", make_summary(item["newTestament"], topics)),
        ]
        for row_idx, (label, value) in enumerate(rows, start=1):
            set_cell_shading(table.cell(row_idx, 0), "F4F4F0")
            set_font(table.cell(row_idx, 0).paragraphs[0].add_run(label), size=9, bold=True)
            set_font(table.cell(row_idx, 1).paragraphs[0].add_run(value), size=9)
        doc.add_paragraph().paragraph_format.space_after = Pt(1)

    doc.add_section(WD_SECTION.NEW_PAGE)
    h = doc.add_paragraph()
    set_font(h.add_run("초안 자체 평가"), size=16, bold=True)
    evaluations = [
        ("완성도", "구조 초안 100%, 본문별 세부 해설 완성도 약 35%. 365개 날짜와 두 본문 칸은 모두 갖추었지만, 현재 문장은 성경 권별 핵심 주제를 적용한 방향성 초안입니다."),
        ("장점", "날짜·본문 누락 없이 한 해 전체를 검토할 수 있고, 하나님의 주권·언약·은혜·그리스도 중심성이라는 개혁주의 방향이 일관됩니다."),
        ("한계", "각 장의 고유 사건과 논증을 충분히 서술하지 못해 같은 성경 권 안에서 문장이 반복됩니다. 사이트 게시용으로 사용하기 전 365일 모두 장별 세부 내용으로 다시 작성해야 합니다."),
        ("신학 검토", "도덕주의나 인간 중심의 적용을 피하는 방향은 적절합니다. 다만 예언서의 역사적 문맥, 율법과 복음의 관계, 시편의 메시아적 적용, 종말론 본문은 개별 검토가 필요합니다."),
        ("권장 수정 순서", "① 날짜와 본문 범위 확인 ② 각 장의 사건·논증을 1–2문장으로 구체화 ③ 그리스도·언약과의 연결이 본문에서 정당한지 검토 ④ 표현 중복 제거 ⑤ 교리·인명·지명 최종 교정"),
        ("게시 판단", "현재 문서는 검토용 골격으로는 사용 가능하지만 홈페이지 공개용 완성 원고로는 부적합합니다."),
    ]
    for label, value in evaluations:
        p = doc.add_paragraph()
        set_font(p.add_run(f"{label}: "), size=10, bold=True)
        set_font(p.add_run(value), size=10)

    doc.save(output_path)
    print(json.dumps({"days": len(plan), "output": str(output_path)}, ensure_ascii=False))


if __name__ == "__main__":
    build(Path(sys.argv[1]), Path(sys.argv[2]), Path(sys.argv[3]))
